"""Le Convoi — extraction de texte brut depuis n'importe quel format de fichier.

Formats supportés (avec dégradation gracieuse si la dépendance optionnelle
n'est pas installée) :
- .txt / .md / .log         → stdlib (utf-8 puis fallback latin-1)
- .csv / .tsv               → csv.DictReader (rendu tabulé)
- .xlsx / .xlsm             → openpyxl si dispo
- .docx                     → python-docx si dispo, sinon unzip + parse XML
- .pdf                      → pypdf si dispo
- .png/.jpg/.jpeg/.bmp/.tif → pytesseract si dispo (OCR), sinon erreur claire

L'extraction renvoie un dict :
    {
        "format": "csv" | "xlsx" | "docx" | "pdf" | "image" | "text",
        "text": "<texte brut concaténé>",
        "rows": [ {col: val}, ... ]   # uniquement pour csv/xlsx, sinon []
        "warnings": [ "..." ],
    }

L'IA prendra ensuite ce dict pour reconstruire la liste structurée.
"""

from __future__ import annotations

import csv
import io
import logging
import re
import zipfile
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

logger = logging.getLogger(__name__)


SUPPORTED_EXTENSIONS = {
    ".txt", ".md", ".log",
    ".csv", ".tsv",
    ".xlsx", ".xlsm",
    ".docx",
    ".pdf",
    ".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp",
}

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp"}


class ParserError(Exception):
    pass


def parse_file(path: str | Path) -> dict[str, Any]:
    """Dispatch vers le bon extracteur selon l'extension."""
    p = Path(path)
    if not p.exists():
        raise ParserError(f"Fichier introuvable : {p}")
    ext = p.suffix.lower()
    if ext in (".txt", ".md", ".log", ""):
        return _parse_text(p)
    if ext == ".csv":
        return _parse_csv(p, delimiter=",")
    if ext == ".tsv":
        return _parse_csv(p, delimiter="\t")
    if ext in (".xlsx", ".xlsm"):
        return _parse_xlsx(p)
    if ext == ".docx":
        return _parse_docx(p)
    if ext == ".pdf":
        return _parse_pdf(p)
    if ext in IMAGE_EXTENSIONS:
        return _parse_image(p)
    raise ParserError(
        f"Format non supporté : {ext}. "
        f"Formats acceptés : {', '.join(sorted(SUPPORTED_EXTENSIONS))}."
    )


# ---------------------------------------------------------------------------
# Implémentations par format
# ---------------------------------------------------------------------------
def _read_bytes(p: Path) -> bytes:
    return p.read_bytes()


def _decode_text(b: bytes) -> str:
    for enc in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return b.decode(enc)
        except UnicodeDecodeError:
            continue
    return b.decode("utf-8", errors="replace")


def _parse_text(p: Path) -> dict[str, Any]:
    text = _decode_text(_read_bytes(p))
    return {"format": "text", "text": text, "rows": [], "warnings": []}


def _parse_csv(p: Path, delimiter: str = ",") -> dict[str, Any]:
    text = _decode_text(_read_bytes(p))
    rows: list[dict[str, str]] = []
    warnings: list[str] = []
    try:
        # Sniff delimiter si possible
        sample = text[:4096]
        try:
            dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
            actual_delim = dialect.delimiter
        except csv.Error:
            actual_delim = delimiter
        reader = csv.DictReader(io.StringIO(text), delimiter=actual_delim)
        for row in reader:
            cleaned = {(k or "").strip(): (v or "").strip() for k, v in row.items()
                       if k is not None}
            if any(cleaned.values()):
                rows.append(cleaned)
    except Exception as exc:
        warnings.append(f"Lecture CSV imparfaite : {exc}")
    return {
        "format": "csv",
        "text": text,
        "rows": rows,
        "warnings": warnings,
    }


def _parse_xlsx(p: Path) -> dict[str, Any]:
    warnings: list[str] = []
    rows: list[dict[str, str]] = []
    text_parts: list[str] = []
    try:
        from openpyxl import load_workbook  # type: ignore
    except ImportError:
        raise ParserError(
            "Lecture Excel : module 'openpyxl' non installé. "
            "Installe-le avec : pip install openpyxl"
        )
    try:
        wb = load_workbook(filename=str(p), data_only=True, read_only=True)
    except Exception as exc:
        raise ParserError(f"Ouverture XLSX impossible : {exc}") from exc

    for sheet in wb.worksheets:
        sheet_rows = list(sheet.iter_rows(values_only=True))
        if not sheet_rows:
            continue
        header = [str(c).strip() if c is not None else "" for c in sheet_rows[0]]
        # Si pas d'en-tête utilisable (que des None), on génère col_1, col_2…
        if not any(header):
            header = [f"col_{i + 1}" for i in range(len(sheet_rows[0]))]
        for raw_row in sheet_rows[1:]:
            cleaned: dict[str, str] = {}
            for i, val in enumerate(raw_row):
                key = header[i] if i < len(header) else f"col_{i + 1}"
                if not key:
                    key = f"col_{i + 1}"
                cleaned[key] = "" if val is None else str(val).strip()
            if any(cleaned.values()):
                rows.append(cleaned)
        # Texte tabulé pour l'IA
        text_parts.append(f"=== Feuille : {sheet.title} ===")
        text_parts.append("\t".join(header))
        for raw_row in sheet_rows[1:]:
            text_parts.append("\t".join(
                "" if v is None else str(v) for v in raw_row
            ))

    wb.close()
    return {
        "format": "xlsx",
        "text": "\n".join(text_parts),
        "rows": rows,
        "warnings": warnings,
    }


def _parse_docx(p: Path) -> dict[str, Any]:
    """Extraction texte via python-docx si dispo, fallback ZIP/XML pur sinon."""
    warnings: list[str] = []
    try:
        from docx import Document  # type: ignore
        doc = Document(str(p))
        parts: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                line = "\t".join(cell.text.strip() for cell in row.cells)
                if line.strip():
                    parts.append(line)
        return {
            "format": "docx",
            "text": "\n".join(parts),
            "rows": [],
            "warnings": warnings,
        }
    except ImportError:
        # Fallback : .docx = ZIP avec word/document.xml
        warnings.append("python-docx absent — extraction XML brute (peut perdre le formatage).")
        try:
            with zipfile.ZipFile(str(p)) as z:
                xml_bytes = z.read("word/document.xml")
            ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
            root = ET.fromstring(xml_bytes)
            parts: list[str] = []
            for para in root.iter(f"{ns}p"):
                texts = [t.text or "" for t in para.iter(f"{ns}t")]
                line = "".join(texts).strip()
                if line:
                    parts.append(line)
            return {
                "format": "docx",
                "text": "\n".join(parts),
                "rows": [],
                "warnings": warnings,
            }
        except Exception as exc:
            raise ParserError(f"Lecture DOCX impossible : {exc}") from exc
    except Exception as exc:
        raise ParserError(f"Lecture DOCX impossible : {exc}") from exc


def _parse_pdf(p: Path) -> dict[str, Any]:
    warnings: list[str] = []
    try:
        import pypdf  # type: ignore
    except ImportError:
        try:
            from PyPDF2 import PdfReader  # type: ignore
        except ImportError:
            raise ParserError(
                "Lecture PDF : module 'pypdf' non installé. "
                "Installe-le avec : pip install pypdf"
            )
        reader_cls = PdfReader
        as_pypdf2 = True
    else:
        reader_cls = pypdf.PdfReader
        as_pypdf2 = False

    try:
        reader = reader_cls(str(p))
        # Selon version : reader.pages
        pages = reader.pages
        parts: list[str] = []
        for i, page in enumerate(pages):
            try:
                txt = page.extract_text() or ""
            except Exception as exc:
                warnings.append(f"Page {i + 1} illisible : {exc}")
                continue
            if txt.strip():
                parts.append(txt)
        text = "\n\n".join(parts)
        if not text.strip():
            warnings.append(
                "PDF sans texte extractible (probablement scanné). "
                "Essaie une conversion OCR au préalable."
            )
        return {
            "format": "pdf",
            "text": text,
            "rows": [],
            "warnings": warnings,
        }
    except Exception as exc:
        raise ParserError(f"Lecture PDF impossible : {exc}") from exc


def _parse_image(p: Path) -> dict[str, Any]:
    """OCR via pytesseract (Tesseract OCR doit être installé sur la machine)."""
    try:
        import pytesseract  # type: ignore
        from PIL import Image  # type: ignore
    except ImportError:
        raise ParserError(
            "OCR image : modules 'pytesseract' + 'Pillow' requis. "
            "Installe-les avec : pip install pytesseract pillow "
            "(et Tesseract OCR doit être installé sur Windows)."
        )
    try:
        img = Image.open(str(p))
        text = pytesseract.image_to_string(img, lang="fra+eng")
        return {
            "format": "image",
            "text": text,
            "rows": [],
            "warnings": [
                "OCR utilisé : la précision dépend de la qualité de l'image. "
                "Vérifie le tableau extrait avant envoi.",
            ],
        }
    except Exception as exc:
        raise ParserError(
            f"OCR a échoué : {exc}. "
            "Vérifie que Tesseract OCR est installé et accessible."
        ) from exc


# ---------------------------------------------------------------------------
# Pré-extraction regex (fallback universel)
# ---------------------------------------------------------------------------
_EMAIL_RE = re.compile(
    r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}",
)
# Téléphone FR + international approximatif (10+ chiffres avec séparateurs)
_PHONE_RE = re.compile(
    r"(?:\+?\d{1,3}[\s.\-]?)?(?:\(?\d{2,4}\)?[\s.\-]?){3,5}\d{2,4}",
)
_URL_RE = re.compile(
    r"\bhttps?://[\w.\-/]+|"
    r"\b(?:www\.)?[a-z0-9\-]+\.(?:fr|com|net|org|eu|io|co|app|shop|store|paris|bzh)\b",
    re.IGNORECASE,
)


# ---------------------------------------------------------------------------
# Mapping direct rows → prospects (fast path pour csv/xlsx déjà structurés)
# ---------------------------------------------------------------------------
# Synonymes de colonnes par champ canonique. Les clés sont normalisées
# (lowercase, sans accents, sans BOM) avant comparaison.
_COLUMN_SYNONYMS: dict[str, tuple[str, ...]] = {
    "email": (
        "email", "e-mail", "mail", "courriel", "adresse mail",
        "adresse email", "email principal", "emails", "address",
    ),
    "raison_sociale": (
        "raison sociale", "entreprise", "societe", "company",
        "nom commercial", "nom entreprise", "structure", "etablissement",
        "organisation", "organization", "categorie google",
    ),
    "prenom": (
        "prenom", "first name", "firstname", "given name",
    ),
    "nom": (
        "nom", "last name", "lastname", "surname", "nom de famille",
        "name", "contact",
    ),
    "telephone": (
        "telephone", "tel", "phone", "mobile", "portable",
        "numero", "tel principal",
    ),
    "site_web": (
        "site web", "site", "website", "url", "site internet",
        "site / reseau", "lien", "web",
    ),
    "adresse": (
        "adresse", "address", "rue", "voie",
    ),
    "ville": (
        "ville", "city", "commune", "localite",
    ),
    "code_postal": (
        "code postal", "cp", "zip", "postal code", "zip code",
    ),
    "secteur": (
        "secteur", "sector", "activite", "type d'activite",
        "metier", "domaine", "industry", "industrie",
        "categorie", "category",
    ),
    "notes": (
        "notes", "note", "commentaire", "comment", "remarques",
        "infos", "description",
    ),
}


def _normalize_key(s: str) -> str:
    """lowercase + strip BOM + retire accents + collapse espaces."""
    if not s:
        return ""
    import unicodedata
    s = s.replace("﻿", "").strip().lower()
    s = unicodedata.normalize("NFKD", s)
    s = "".join(c for c in s if not unicodedata.combining(c))
    return " ".join(s.split())


def _build_column_index(headers: list[str]) -> dict[str, str]:
    """Renvoie {champ_canonique: header_original} pour les colonnes reconnues.

    Si plusieurs colonnes matchent un même champ, on prend la première qui
    apparaît dans le fichier (ordre de la liste headers).
    """
    norm_headers = [(h, _normalize_key(h)) for h in headers]
    out: dict[str, str] = {}
    for field, synonyms in _COLUMN_SYNONYMS.items():
        for raw_h, norm_h in norm_headers:
            if not norm_h:
                continue
            # Match exact d'abord
            if norm_h in synonyms:
                out[field] = raw_h
                break
        if field in out:
            continue
        # Fallback : substring (ex: "email principal" contient "email")
        for raw_h, norm_h in norm_headers:
            if not norm_h:
                continue
            if any(syn in norm_h for syn in synonyms):
                out[field] = raw_h
                break
    return out


def rows_to_prospects(rows: list[dict[str, str]]) -> list[dict[str, str]]:
    """Mappe directement des lignes csv/xlsx vers des prospects structurés.

    Renvoie [] si aucune colonne email reconnaissable n'est trouvée — dans
    ce cas l'appelant retombera sur l'extraction IA.
    """
    if not rows:
        return []
    headers = list(rows[0].keys())
    col = _build_column_index(headers)
    if "email" not in col:
        return []

    fields = [
        "raison_sociale", "prenom", "nom", "email", "telephone",
        "site_web", "adresse", "ville", "code_postal", "secteur", "notes",
    ]
    out: list[dict[str, str]] = []
    for row in rows:
        prospect = {f: "" for f in fields}
        for field, header in col.items():
            v = row.get(header, "")
            prospect[field] = (str(v).strip() if v is not None else "")
        # Validation minimale : on garde la ligne si elle a un email plausible
        email = prospect.get("email", "")
        if "@" in email and "." in email.split("@")[-1]:
            out.append(prospect)
    return out


def harvest_basics(text: str) -> dict[str, list[str]]:
    """Extrait emails, téléphones, URLs depuis du texte brut. Universel.

    Sert de complément à l'IA : si l'IA rate des contacts, on peut au moins
    les présenter manuellement à l'utilisateur.
    """
    emails = sorted(set(m.group(0).lower() for m in _EMAIL_RE.finditer(text)))
    urls = sorted(set(m.group(0).lower() for m in _URL_RE.finditer(text)))
    # Téléphones : on filtre en post car la regex est lousy
    raw_phones = [m.group(0) for m in _PHONE_RE.finditer(text)]
    phones: list[str] = []
    seen: set[str] = set()
    for raw in raw_phones:
        digits = re.sub(r"\D", "", raw)
        if 9 <= len(digits) <= 15 and digits not in seen:
            seen.add(digits)
            phones.append(raw.strip())
    return {"emails": emails, "phones": phones, "urls": urls}
